"""Export project pages to JSON / TXT / DOCX."""
import json
import os
import re
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_DIRECTION

# Categories rendered as flowing text (everything else is skipped or handled separately)
TEXT_CATEGORIES = {'Text', 'Section-header', 'Title', 'Caption', 'Footnote', 'List-item'}
SKIP_CATEGORIES = {'Picture', 'Page-header', 'Page-footer', 'Page-number', 'Formula'}



# ══════════════════════════════════════════════════════════════════════
# TEXT MODE: block content can contain leftover HTML tags (from the
# contenteditable review editor) and/or markdown syntax (from the original
# OCR text). "raw" export keeps that verbatim (old behavior). "formatted"
# (the new default) turns it into clean, human-readable output: for TXT
# that means tags/markdown stripped; for DOCX it means real bold/italic/
# underline/superscript Word formatting instead of visible tag characters.
# ══════════════════════════════════════════════════════════════════════

_BLOCK_BOUNDARY_RE = re.compile(r'</\s*(?:p|div|li|h[1-6]|tr)\s*>', re.IGNORECASE)
_TAG_RE = re.compile(r'<[^>]+>')
_MD_HEADER_RE = re.compile(r'^\s{0,3}#{1,6}\s*', re.MULTILINE)
_MD_BOLD_RE = re.compile(r'\*\*(.+?)\*\*|__(.+?)__')
_MD_ITALIC_RE = re.compile(r'(?<!\*)\*(?!\*)(.+?)\*(?!\*)|(?<!_)_(?!_)(.+?)_(?!_)')

_ARABIC_CHAR_RE = re.compile(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]')
_LATIN_CHAR_RE = re.compile(r'[a-zA-Z]')

def _detect_text_direction(text, default_rtl=True):
    """Auto-detect paragraph text direction if not explicitly set in HTML/element metadata."""
    if not text:
        return default_rtl
    clean = _TAG_RE.sub('', text)
    num_arabic = len(_ARABIC_CHAR_RE.findall(clean))
    num_latin = len(_LATIN_CHAR_RE.findall(clean))
    if num_latin > 0 and num_arabic == 0:
        return False  # LTR for English/Latin text
    elif num_arabic > 0:
        return True   # RTL for Arabic text
    return default_rtl

def _extract_html_dir(text):
    """Extract direction (rtl/ltr) from root or inline HTML tags (e.g. dir="rtl" or style="direction: rtl")."""
    if not text:
        return None
    m = re.search(r'\bdir\s*=\s*["\']?(rtl|ltr)["\']?', text, re.IGNORECASE)
    if m:
        return m.group(1).lower()
    m2 = re.search(r'direction\s*:\s*(rtl|ltr)', text, re.IGNORECASE)
    if m2:
        return m2.group(1).lower()
    return None

def _extract_html_align(text):
    """Extract block alignment from root or inline HTML tags (e.g. style="text-align: center" or align="center")."""
    if not text:
        return None
    m = re.search(r'text-align\s*:\s*(right|left|center|justify|both|start|end|lowkashida|mediumkashida|highkashida|distribute)', text, re.IGNORECASE)
    if m:
        return m.group(1).lower()
    m2 = re.search(r'\balign\s*=\s*["\']?(right|left|center|justify|both|start|end)["\']?', text, re.IGNORECASE)
    if m2:
        return m2.group(1).lower()
    return None

def _extract_html_line_spacing(text):
    """Extract line-height from inline HTML tags (e.g. style="line-height: 1.5")."""
    if not text:
        return None
    m = re.search(r'line-height\s*:\s*([\d.]+)', text, re.IGNORECASE)
    if m:
        try:
            return float(m.group(1))
        except ValueError:
            pass
    return None

def _extract_html_margin_bottom(text):
    """Extract margin-bottom from inline HTML tags (e.g. style="margin-bottom: 12pt")."""
    if not text:
        return None
    m = re.search(r'margin-bottom\s*:\s*([\d.]+)pt', text, re.IGNORECASE)
    if m:
        try:
            return float(m.group(1))
        except ValueError:
            pass
    return None

def _extract_html_margin_top(text):
    """Extract margin-top from inline HTML tags (e.g. style="margin-top: 12pt")."""
    if not text:
        return None
    m = re.search(r'margin-top\s*:\s*([\d.]+)pt', text, re.IGNORECASE)
    if m:
        try:
            return float(m.group(1))
        except ValueError:
            pass
    return None

def _strip_markdown_and_tags(segment):
    """Best-effort: drop leftover HTML tags, but preserve <br> as newlines."""
    # 1. Convert raw OCR soft returns (\n) into spaces, exactly as the browser renders them.
    s = segment.replace('\n', ' ').replace('\r', '')
    
    # 2. Convert explicit HTML line breaks (from user edits in the UI) into real text newlines.
    s = re.sub(r'(?i)<br\s*/?>', '\n', s)
    
    # 3. Strip all remaining HTML tags (like <p>, <span>)
    s = _TAG_RE.sub('', s)
    s = _MD_HEADER_RE.sub('', s.strip())
    s = _MD_BOLD_RE.sub(lambda m: m.group(1) or m.group(2) or '', s)
    s = _MD_ITALIC_RE.sub(lambda m: m.group(1) or m.group(2) or '', s)
    
    # 4. Clean up any accidental double spaces created by step 1
    s = s.replace('  ', ' ')
    
    return s.strip()
    
def _split_block_paragraphs(text):
    """Split raw (HTML/markdown) text into paragraph-level chunks on block
    boundaries (</p>, </div>, </li>, etc.)."""
    if not text:
        return []
    parts = _BLOCK_BOUNDARY_RE.split(text)
    return [p.strip() for p in parts if p and p.strip()]


def format_display_text(text):
    """Formatted/processed plain text: tags and markdown syntax removed,
    block breaks preserved as newlines. Used for TXT export and text preview."""
    cleaned = [_strip_markdown_and_tags(p) for p in _split_block_paragraphs(text)]
    return '\n'.join(c for c in cleaned if c)


# ---- Inline run parsing (used by DOCX formatted export) ----
_MD_INLINE_TOKEN_RE = re.compile(r'(\*\*|__|\*|_)')

_BOLD_TAGS = {'b', 'strong'}
_ITALIC_TAGS = {'i', 'em'}
_UNDERLINE_TAGS = {'u'}
_SUP_TAGS = {'sup'}
_SUB_TAGS = {'sub'}

_CSS_BOLD_RE = re.compile(r'font-weight\s*:\s*(bold|bolder|[6-9]00)', re.IGNORECASE)
_CSS_ITALIC_RE = re.compile(r'font-style\s*:\s*italic', re.IGNORECASE)
_CSS_UNDERLINE_RE = re.compile(r'text-decoration(?:-line)?\s*:\s*[^;]*underline', re.IGNORECASE)
_CSS_STRIKE_RE = re.compile(r'text-decoration(?:-line)?\s*:\s*[^;]*line-through', re.IGNORECASE)
_CSS_BG_COLOR_RE = re.compile(r'background-color\s*:\s*([^;]+)', re.IGNORECASE)
_CSS_FG_COLOR_RE = re.compile(r'(?<![-\w])color\s*:\s*([^;]+)', re.IGNORECASE)

_STRIKE_TAGS = {'strike', 's', 'del'}

_NAMED_COLORS = {
    'red': 'FF0000', 'green': '008000', 'blue': '0000FF', 'yellow': 'FFFF00',
    'black': '000000', 'white': 'FFFFFF', 'orange': 'FFA500', 'purple': '800080',
    'gray': '808080', 'grey': '808080', 'pink': 'FFC0CB', 'brown': 'A52A2A',
    'cyan': '00FFFF', 'magenta': 'FF00FF',
}


def _css_color_to_hex(value):
    """Convert a CSS color value into a 6-char uppercase hex string, or None if unrecognized."""
    if not value:
        return None
    value = str(value).strip().lower()
    if value.startswith('#'):
        value = value[1:]
    if len(value) == 3 and all(c in '0123456789abcdef' for c in value):
        return ''.join(c * 2 for c in value).upper()
    if len(value) >= 6 and all(c in '0123456789abcdef' for c in value[:6]):
        return value[:6].upper()
    m = re.match(r'rgba?\(([^)]+)\)', value)
    if m:
        parts = [p.strip() for p in m.group(1).split(',')]
        try:
            r, g, b = (max(0, min(255, int(float(parts[i])))) for i in range(3))
            return f'{r:02X}{g:02X}{b:02X}'
        except Exception:
            return None
    return _NAMED_COLORS.get(value)


def _extract_css_colors(css):
    """Returns (text_color_hex, highlight_color_hex) from a style="..." string."""
    if not css:
        return None, None
    bg_match = _CSS_BG_COLOR_RE.search(css)
    bg_hex = _css_color_to_hex(bg_match.group(1)) if bg_match else None
    css_wo_bg = _CSS_BG_COLOR_RE.sub('', css)
    fg_match = _CSS_FG_COLOR_RE.search(css_wo_bg)
    fg_hex = _css_color_to_hex(fg_match.group(1)) if fg_match else None
    return fg_hex, bg_hex


def _style_from_tag(tag, attrs_dict):
    """Determine bold/italic/underline/strike/superscript/subscript/color/
    highlight/font_family/font_size/dir/align from a tag name AND from any inline style="..."
    attribute (or legacy <font face="..." size="..." dir="...">)."""
    style = {'bold': False, 'italic': False, 'underline': False, 'strike': False,
             'superscript': False, 'subscript': False, 'color': None, 'highlight': None,
             'font_family': None, 'font_size': None, 'dir': None, 'align': None}

    dir_attr = attrs_dict.get('dir')
    if dir_attr:
        style['dir'] = dir_attr.strip().lower()

    align_attr = attrs_dict.get('align')
    if align_attr:
        style['align'] = align_attr.strip().lower()

    if tag in _BOLD_TAGS:
        style['bold'] = True
    if tag in _ITALIC_TAGS:
        style['italic'] = True
    if tag in _UNDERLINE_TAGS:
        style['underline'] = True
    if tag in _STRIKE_TAGS:
        style['strike'] = True
    if tag in _SUP_TAGS:
        style['superscript'] = True
    if tag in _SUB_TAGS:
        style['subscript'] = True
    if tag == 'font':
        color_attr = attrs_dict.get('color')
        if color_attr:
            style['color'] = _css_color_to_hex(color_attr) or _NAMED_COLORS.get(color_attr.strip().lower())
        face_attr = attrs_dict.get('face')
        if face_attr:
            style['font_family'] = face_attr.strip().strip("'\"")
        size_attr = attrs_dict.get('size')
        if size_attr:
            m = re.search(r'\d+', size_attr)
            if m:
                html_sizes = {1: 10, 2: 13, 3: 16, 4: 18, 5: 24, 6: 32, 7: 48}
                style['font_size'] = html_sizes.get(int(m.group(0)), 16)

    css = attrs_dict.get('style') or ''
    if css:
        if _CSS_BOLD_RE.search(css):
            style['bold'] = True
        if _CSS_ITALIC_RE.search(css):
            style['italic'] = True
        if _CSS_UNDERLINE_RE.search(css):
            style['underline'] = True
        if _CSS_STRIKE_RE.search(css):
            style['strike'] = True

        m_dir = re.search(r'direction\s*:\s*(rtl|ltr)', css, re.IGNORECASE)
        if m_dir:
            style['dir'] = m_dir.group(1).lower()

        m_align = re.search(r'text-align\s*:\s*(right|left|center|justify|both|start|end|lowkashida|mediumkashida|highkashida|distribute)', css, re.IGNORECASE)
        if m_align:
            style['align'] = m_align.group(1).lower()

        m_fam = re.search(r'font-family\s*:\s*([^;]+)', css, re.IGNORECASE)
        if m_fam:
            style['font_family'] = m_fam.group(1).strip().strip("'\"")

        m_sz = re.search(r'font-size\s*:\s*([^;]+)', css, re.IGNORECASE)
        if m_sz:
            m_pt = re.search(r'(\d+(?:\.\d+)?)\s*(pt|px|em)?', m_sz.group(1), re.IGNORECASE)
            if m_pt:
                num = float(m_pt.group(1))
                unit = (m_pt.group(2) or '').lower()
                if unit == 'px':
                    num = num * 0.75
                elif unit == 'em':
                    num = num * 12.0
                style['font_size'] = int(round(num))

        fg_hex, bg_hex = _extract_css_colors(css)
        if fg_hex:
            style['color'] = fg_hex
        if bg_hex:
            style['highlight'] = bg_hex
    return style


class _InlineRunParser(HTMLParser):
    """Walks a fragment of HTML and produces a flat list of (text, style)
    runs, tracking formatting via a tag stack so nested/attributed tags are all recognized."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.runs = []
        self._stack = []  # list of style dicts, one per currently-open tag
        self._md_state = {'bold': False, 'italic': False}

    def _current_style(self):
        merged = {'bold': False, 'italic': False, 'underline': False, 'strike': False,
                  'superscript': False, 'subscript': False, 'color': None, 'highlight': None,
                  'font_family': None, 'font_size': None, 'dir': None, 'align': None}
        for s in self._stack:
            for k in ('bold', 'italic', 'underline', 'strike', 'superscript', 'subscript'):
                if s.get(k):
                    merged[k] = True
            if s.get('color'):
                merged['color'] = s['color']
            if s.get('highlight'):
                merged['highlight'] = s['highlight']
            if s.get('font_family'):
                merged['font_family'] = s['font_family']
            if s.get('font_size'):
                merged['font_size'] = s['font_size']
            if s.get('dir'):
                merged['dir'] = s['dir']
            if s.get('align'):
                merged['align'] = s['align']
        if self._md_state['bold']:
            merged['bold'] = True
        if self._md_state['italic']:
            merged['italic'] = True
        return merged

    def handle_starttag(self, tag, attrs):
        if tag == 'br':
            self.runs.append(('\n', {}))
            return
        self._stack.append(_style_from_tag(tag, dict(attrs)))

    def handle_startendtag(self, tag, attrs):
        if tag == 'br':
            self.runs.append(('\n', {}))

    def handle_endtag(self, tag):
        if self._stack:
            self._stack.pop()

    def handle_data(self, data):
        if not data:
            return
        for tok in (t for t in _MD_INLINE_TOKEN_RE.split(data) if t != ''):
            if tok in ('**', '__'):
                self._md_state['bold'] = not self._md_state['bold']
                continue
            if tok in ('*', '_'):
                self._md_state['italic'] = not self._md_state['italic']
                continue
            self.runs.append((tok, self._current_style()))


def parse_inline_runs(text):
    """Convert a paragraph fragment of HTML/markdown-ish text into a list of (plain_text, style) runs."""
    if not text:
        return []
    text = _MD_HEADER_RE.sub('', text)
    text = text.replace('\n', ' ').replace('\r', '')
    text = text.replace('  ', ' ')
    
    parser = _InlineRunParser()
    try:
        parser.feed(text)
        parser.close()
    except Exception:
        cleaned = _strip_markdown_and_tags(text)
        return [(cleaned, {})] if cleaned else []
        
    runs = []
    for run_text, style in parser.runs:
        if run_text == '\n' or run_text.strip() or run_text == ' ':
            runs.append((run_text, style))
    return runs

def _set_paragraph_spacing(para, space_before=0, space_after=6):
    fmt = para.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)


def _set_line_spacing(para, lines=1.0):
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = lines


def _add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(break_type=WD_BREAK.PAGE)
    _set_paragraph_spacing(para, 0, 0)


def _add_page_number_label(doc, label, font_name, font_size, rtl):
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, 0, 0)
    _apply_paragraph_layout(para, 'center', rtl)
    run = para.add_run(label)
    _set_run_font_and_bidi(run, font_name, font_size, is_rtl=rtl)
    run.italic = True


def _set_section_rtl(section, is_rtl=True):
    """Set section-level BiDi direction tag (w:bidi) so Word renders page layout in RTL."""
    sectPr = section._sectPr
    bidi = sectPr.find(qn('w:bidi'))
    if is_rtl:
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            sectPr.append(bidi)
        bidi.set(qn('w:val'), '1')
    else:
        if bidi is not None:
            sectPr.remove(bidi)


def _set_style_rtl(style, font_name=None, font_size_pt=None, is_rtl=True):
    """Set Normal style paragraph & run properties for BiDi/RTL."""
    pPr = style._element.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if is_rtl:
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            pPr.insert(0, bidi)
        bidi.set(qn('w:val'), '1')
    else:
        if bidi is not None:
            pPr.remove(bidi)
    _reorder_pPr(pPr)

    rPr = style._element.get_or_add_rPr()
    if font_name:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        if is_rtl:
            rFonts.set(qn('w:hint'), 'cs')

    if font_size_pt:
        half_pts = str(int(round(float(font_size_pt) * 2)))
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rPr.append(sz)
        sz.set(qn('w:val'), half_pts)

        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), half_pts)

    rtl_elem = rPr.find(qn('w:rtl'))
    if is_rtl:
        if rtl_elem is None:
            rtl_elem = OxmlElement('w:rtl')
            rPr.append(rtl_elem)
        rtl_elem.set(qn('w:val'), '1')
        
        cs_elem = rPr.find(qn('w:cs'))
        if cs_elem is None:
            cs_elem = OxmlElement('w:cs')
            rPr.append(cs_elem)
        cs_elem.set(qn('w:val'), '1')
    else:
        if rtl_elem is not None:
            rPr.remove(rtl_elem)
        cs_elem = rPr.find(qn('w:cs'))
        if cs_elem is not None:
            rPr.remove(cs_elem)

    _reorder_rPr(rPr)


def _set_run_font_and_bidi(run, font_name=None, font_size_pt=None, is_rtl=True):
    """
    Sets run-level font family, font size, and RTL direction in OOXML (w:rPr),
    populating BOTH Latin (w:ascii, w:hAnsi, w:sz) AND Complex Script / Arabic
    (w:cs, w:szCs, w:rtl, w:cs, w:hint="cs") so Microsoft Word correctly renders Arabic fonts, sizes,
    and right-to-left layout without resetting to Word default sizes/fonts or LTR direction.
    """
    rPr = run._element.get_or_add_rPr()

    # 1. Font Family (Latin + Complex Script / Arabic + Complex Script hint)
    if font_name:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        if is_rtl:
            rFonts.set(qn('w:hint'), 'cs')

    # 2. Font Size (Latin + Complex Script / Arabic)
    if font_size_pt:
        half_pts = str(int(round(float(font_size_pt) * 2)))
        
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rPr.append(sz)
        sz.set(qn('w:val'), half_pts)

        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), half_pts)

    # 3. Run-level RTL (w:rtl) & Complex Script (w:cs)
    if is_rtl:
        rtl_elem = rPr.find(qn('w:rtl'))
        if rtl_elem is None:
            rtl_elem = OxmlElement('w:rtl')
            rPr.append(rtl_elem)
        rtl_elem.set(qn('w:val'), '1')
        
        cs_elem = rPr.find(qn('w:cs'))
        if cs_elem is None:
            cs_elem = OxmlElement('w:cs')
            rPr.append(cs_elem)
        cs_elem.set(qn('w:val'), '1')
    else:
        rtl_elem = rPr.find(qn('w:rtl'))
        if rtl_elem is not None:
            rtl_elem.set(qn('w:val'), '0')
        cs_elem = rPr.find(qn('w:cs'))
        if cs_elem is not None:
            cs_elem.set(qn('w:val'), '0')

    _reorder_rPr(rPr)


def _apply_run_formatting(run, style, font_name, font_size, is_r_rtl=True, cat_fmt=None):
    """Applies complete inline formatting to a docx run, including Complex Script font, size,
    bidi, bold (w:b & w:bCs), italic (w:i & w:iCs), underline, strike, color (w:color),
    and highlight shading (w:shd), then guarantees ISO/IEC 29500 element sequence order via _reorder_rPr."""
    cat_fmt = cat_fmt or {}
    rPr = run._element.get_or_add_rPr()

    r_font_name = style.get('font_family') or font_name
    r_font_size = style.get('font_size') or font_size

    _set_run_font_and_bidi(run, r_font_name, r_font_size, is_rtl=is_r_rtl)

    # Bold (Standard + Complex Script)
    if style.get('bold') or cat_fmt.get('bold'):
        run.bold = True
        bCs = rPr.find(qn('w:bCs'))
        if bCs is None:
            bCs = OxmlElement('w:bCs')
            rPr.append(bCs)
        bCs.set(qn('w:val'), '1')

    # Italic (Standard + Complex Script)
    if style.get('italic') or cat_fmt.get('italic'):
        run.italic = True
        iCs = rPr.find(qn('w:iCs'))
        if iCs is None:
            iCs = OxmlElement('w:iCs')
            rPr.append(iCs)
        iCs.set(qn('w:val'), '1')

    # Underline
    if style.get('underline') or cat_fmt.get('underline'):
        run.underline = True

    # Strike
    if style.get('strike'):
        run.font.strike = True

    # Superscript / Subscript
    if style.get('superscript'):
        run.font.superscript = True
    elif style.get('subscript'):
        run.font.subscript = True

    # Text Color (run.font.color + explicit w:color OOXML element)
    run_color = style.get('color') or cat_fmt.get('color')
    if run_color:
        try:
            hex_c = _css_color_to_hex(run_color)
            if hex_c:
                run.font.color.rgb = RGBColor.from_string(hex_c)
                col_el = rPr.find(qn('w:color'))
                if col_el is None:
                    col_el = OxmlElement('w:color')
                    rPr.append(col_el)
                col_el.set(qn('w:val'), hex_c)
        except Exception:
            pass

    # Shading / Highlight
    run_bg = style.get('highlight') or cat_fmt.get('bgColor')
    if run_bg:
        try:
            hex_bg = _css_color_to_hex(run_bg) or run_bg
            if hex_bg:
                _set_run_highlight_hex(run, hex_bg)
        except Exception:
            pass

    # Strictly reorder w:rPr child tags to match OpenXML standard
    _reorder_rPr(rPr)



# ══════════════════════════════════════════════════════════════════════
# ARABIC POETRY TABLE HELPERS
# Handles Vertical-poetry (classical 2-hemistich per row) and
# Staggered-poetry (staggered: hemistichs on alternating rows)
# ══════════════════════════════════════════════════════════════════════

def _set_table_no_borders(tbl):
    """Remove all visible borders from a Word table via raw OOXML."""
    tblPr = tbl._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._element.insert(0, tblPr)
    # Remove old borders element if present
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _set_row_height(row, height_twips, exact=True):
    """Set row height in twips (20 twips = 1 pt) via OOXML w:trHeight.
    If exact=True, uses hRule='exact' so the row height is strictly fixed,
    hiding any trailing soft-return line break."""
    trPr = row._tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        row._tr.insert(0, trPr)
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is None:
        trHeight = OxmlElement('w:trHeight')
        trPr.append(trHeight)
    trHeight.set(qn('w:val'), str(int(height_twips)))
    trHeight.set(qn('w:hRule'), 'exact' if exact else 'atLeast')


def _set_cell_width(cell, width_cm):
    """Set explicit cell width in cm for OOXML table layout."""
    width_twips = int(width_cm * 567)  # 1 cm = 567 dxa/twips
    cell.width = Cm(width_cm)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')


def _set_cell_valign(cell, valign='bottom'):
    """Set vertical alignment (top/center/bottom) of a Word table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        tcPr.append(vAlign)
    vAlign.set(qn('w:val'), valign)

def _apply_poetry_paragraph_layout_v169(para, align, is_rtl):
    """Restore exact Version 169 paragraph layout logic exclusively for poetry cells."""
    pPr = para._p.get_or_add_pPr()
    
    # 1. BiDi direction tag MUST be first inside pPr for Word to render RTL table paragraphs properly
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.insert(0, bidi)
    bidi.set(qn('w:val'), '1' if is_rtl else '0')

    # 2. Alignment tag
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), align)

def _fill_poetry_cell(cell, text, font_name, font_size, is_rtl, align='lowKashida',
                      valign='bottom', soft_return=True, cat_fmt=None):
    """Write text into a poetry table cell with proper font, size, direction,
    Justify Low (w:jc=lowKashida), vertical bottom alignment, and soft return (Shift+Enter)."""
    cell.text = ''
    para = cell.paragraphs[0]
    _set_paragraph_spacing(para, 0, 0)
    _apply_poetry_paragraph_layout_v169(para, align, is_rtl)
    
    # Vertical alignment: bottom
    _set_cell_valign(cell, valign)
    
    if text and text.strip():
        runs_spec = parse_inline_runs(text)
        for run_text, style in runs_spec:
            if not run_text:
                continue
            run = para.add_run(run_text)
            _apply_run_formatting(run, style, font_name, font_size, is_r_rtl=is_rtl, cat_fmt=cat_fmt)
            
        # Soft return (Shift+Enter) forces Word to stretch hemistich across full cell width
        if soft_return:
            br_run = para.add_run()
            br_elem = OxmlElement('w:br')
            br_run._element.append(br_elem)
            _set_run_font_and_bidi(br_run, font_name, font_size, is_rtl=is_rtl)
            
    # Re-apply alignment after adding runs
    _apply_poetry_paragraph_layout_v169(para, align, is_rtl)


def _parse_poetry_lines(el):
    """Parse verse lines from a poetry OCR element.
    Returns list of (right_hemistich, left_hemistich) string tuples.
    Prefers table_structure if present, otherwise splits el['text'] on <br>/newline and '|'.
    Handles both 2-column (col0=صدر, col1=عجز) and 3-column
    (col0=صدر, col1=separator, col2=عجز) table_structure formats."""
    lines = []

    # --- Method 1: table_structure cells ---
    ts = el.get('table_structure')
    if ts:
        num_rows = ts.get('rows', 0)
        num_cols = ts.get('cols', 2)
        cells = ts.get('cells', [])
        if num_rows > 0 and cells:
            grid = [[None] * max(num_cols, 2) for _ in range(num_rows)]
            for c_info in cells:
                r, c = c_info.get('row', 0), c_info.get('col', 0)
                if r < num_rows and c < max(num_cols, 2):
                    raw_cell = c_info.get('text', '')
                    grid[r][c] = _strip_markdown_and_tags(raw_cell).strip()
            for row in grid:
                if num_cols >= 3:
                    # 3-column: col0=صدر (right), col1=separator, col2=عجز (left)
                    right = (row[0] or '').strip()
                    left  = (row[2] or '').strip() if len(row) > 2 else ''
                else:
                    # 2-column: col0=صدر (right), col1=عجز (left)
                    right = (row[0] or '').strip()
                    left  = (row[1] or '').strip() if len(row) > 1 else ''
                if right or left:
                    lines.append((right, left))
            if lines:
                return lines

    # --- Method 2: text field split on <br> and '|' ---
    raw = el.get('text', '')
    temp_lines = []
    # Split on <br> tags
    verse_lines = re.split(r'<br\s*/?>', raw, flags=re.IGNORECASE)
    if len(verse_lines) <= 1:
        verse_lines = raw.split('\n')
    for verse in verse_lines:
        verse_clean = _strip_markdown_and_tags(verse).strip()
        if not verse_clean:
            continue
        if '|' in verse_clean:
            parts = verse_clean.split('|', 1)
            temp_lines.append((parts[0].strip(), parts[1].strip()))
        else:
            temp_lines.append((verse_clean, None))
            
    i = 0
    while i < len(temp_lines):
        r, l = temp_lines[i]
        if l is not None:
            lines.append((r, l))
            i += 1
        else:
            if i + 1 < len(temp_lines) and temp_lines[i+1][1] is None:
                lines.append((r, temp_lines[i+1][0]))
                i += 2
            else:
                lines.append((r, ''))
                i += 1
                
    return lines


def _add_poetry_docx(doc, el, cat, font_name, font_size, effective_rtl, cat_fmt=None):
    """Add a poetry block as a borderless Word table with proportional columns,
    exact row heights, vertical bottom alignment, and soft-return justification.

    Vertical-poetry  — 3-column proportional layout:
        [صدر البيت (7.2cm, 45%)] | [المسافة الفاصلة (1.6cm, 10%)] | [عجز البيت (7.2cm, 45%)]
        Exact fixed row height hides soft return; text is justified across full cell width.

    Staggered-poetry  — 2-column staggered layout:
        Row 1: [صدر البيت (7.5cm)] | [empty (7.5cm)]
        Row 2: [empty (7.5cm)]     | [عجز البيت (7.5cm)]
    """
    cat_fmt = cat_fmt or {}
    eff_font_name = cat_fmt.get('fontFamily') or font_name or 'Simplified Arabic'
    eff_font_size = font_size or 14
    if cat_fmt.get('fontSize'):
        m = re.search(r'\d+', str(cat_fmt['fontSize']))
        if m:
            eff_font_size = int(m.group(0))

    lines = _parse_poetry_lines(el)
    if not lines:
        return

    # Exact row height: font_size × 1.6 × 20 twips/pt
    # (exact=True clamps height so the soft-return line break is hidden below cell boundary)
    row_height_twips = int(eff_font_size * 1.6 * 20)

    if cat == 'Vertical-poetry':
        # 3-column table with proportionate widths: 7.2 cm (45%), 1.6 cm (10%), 7.2 cm (45%)
        col_widths = [7.2, 1.6, 7.2]
        tbl = doc.add_table(rows=len(lines), cols=3)
        _set_table_no_borders(tbl)
        if effective_rtl:
            tbl.table_direction = WD_TABLE_DIRECTION.RTL
        for row_idx, (right_text, left_text) in enumerate(lines):
            row = tbl.rows[row_idx]
            _set_row_height(row, row_height_twips, exact=True)
            # Set explicit cell widths
            for c_idx, w in enumerate(col_widths):
                _set_cell_width(row.cells[c_idx], w)

            # In RTL tables: col0 = rightmost (صدر), col1 = center, col2 = leftmost (عجز)
            _fill_poetry_cell(row.cells[0], right_text, eff_font_name, eff_font_size,
                               effective_rtl, align='lowKashida', valign='bottom', soft_return=True, cat_fmt=cat_fmt)
            _fill_poetry_cell(row.cells[1], '',         eff_font_name, eff_font_size,
                               effective_rtl, align='center',     valign='bottom', soft_return=False, cat_fmt=cat_fmt)
            _fill_poetry_cell(row.cells[2], left_text,  eff_font_name, eff_font_size,
                               effective_rtl, align='lowKashida', valign='bottom', soft_return=True, cat_fmt=cat_fmt)

    elif cat == 'Staggered-poetry':
        # 2-column staggered layout: 7.5 cm / 7.5 cm
        col_widths = [7.5, 7.5]
        num_rows = len(lines) * 2
        tbl = doc.add_table(rows=num_rows, cols=2)
        _set_table_no_borders(tbl)
        if effective_rtl:
            tbl.table_direction = WD_TABLE_DIRECTION.RTL
        for line_idx, (right_text, left_text) in enumerate(lines):
            # Odd row: right hemistich (col 0), left col empty
            right_row = tbl.rows[line_idx * 2]
            _set_row_height(right_row, row_height_twips, exact=True)
            for c_idx, w in enumerate(col_widths):
                _set_cell_width(right_row.cells[c_idx], w)
            _fill_poetry_cell(right_row.cells[0], right_text, eff_font_name, eff_font_size,
                               effective_rtl, align='lowKashida', valign='bottom', soft_return=True, cat_fmt=cat_fmt)
            _fill_poetry_cell(right_row.cells[1], '',         eff_font_name, eff_font_size,
                               effective_rtl, align='center',     valign='bottom', soft_return=False, cat_fmt=cat_fmt)

            # Even row: right col empty, left hemistich (col 1)
            left_row = tbl.rows[line_idx * 2 + 1]
            _set_row_height(left_row, row_height_twips, exact=True)
            for c_idx, w in enumerate(col_widths):
                _set_cell_width(left_row.cells[c_idx], w)
            _fill_poetry_cell(left_row.cells[0], '',        eff_font_name, eff_font_size,
                               effective_rtl, align='center',     valign='bottom', soft_return=False, cat_fmt=cat_fmt)
            _fill_poetry_cell(left_row.cells[1], left_text, eff_font_name, eff_font_size,
                               effective_rtl, align='lowKashida', valign='bottom', soft_return=True, cat_fmt=cat_fmt)


def _set_run_highlight_hex(run, hex_color):
    """Use raw OOXML shading (w:shd) on the run so exact custom colors are preserved."""
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    rpr.append(shd)


_PYTHON_DOCX_ALIGN_MAP = {
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    'both': WD_ALIGN_PARAGRAPH.JUSTIFY,
    'lowkashida': WD_ALIGN_PARAGRAPH.JUSTIFY_LOW,
    'lowKashida': WD_ALIGN_PARAGRAPH.JUSTIFY_LOW,
    'mediumkashida': WD_ALIGN_PARAGRAPH.JUSTIFY_HI,
    'mediumKashida': WD_ALIGN_PARAGRAPH.JUSTIFY_HI,
    'highkashida': WD_ALIGN_PARAGRAPH.JUSTIFY_HI,
    'highKashida': WD_ALIGN_PARAGRAPH.JUSTIFY_HI,
    'distribute': WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}

_OPENXML_JC_MAP = {
    'right': 'right',
    'left': 'left',
    'center': 'center',
    'justify': 'both',
    'both': 'both',
    'lowKashida': 'lowKashida',
    'lowkashida': 'lowKashida',
    'mediumKashida': 'mediumKashida',
    'mediumkashida': 'mediumKashida',
    'highKashida': 'highKashida',
    'highkashida': 'highKashida',
    'distribute': 'distribute',
}

_PPR_CHILD_ORDER = [
    'pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr',
    'numPr', 'suppressLineNumbers', 'pBdr', 'shd', 'tabs',
    'suppressAutoHyphens', 'kinsoku', 'wordWrap', 'overflowPunct', 'topLinePunct',
    'autoSpaceDE', 'autoSpaceDN', 'bidi', 'adjustRightInd', 'snapToGrid',
    'spacing', 'ind', 'contextualSpacing', 'mirrorIndents', 'suppressOverlap',
    'jc', 'textDirection', 'textAlignment', 'textboxTightWrap', 'outlineLvl',
    'divId', 'cnfStyle', 'rPr'
]

_PPR_ORDER_INDEX = {tag: i for i, tag in enumerate(_PPR_CHILD_ORDER)}

_RPR_CHILD_ORDER = [
    'rStyle', 'rFonts', 'b', 'bCs', 'i', 'iCs', 'caps', 'smallCaps',
    'strike', 'dstrike', 'outline', 'shadow', 'emboss', 'imprint',
    'noProof', 'snapToGrid', 'vanish', 'webHidden', 'color', 'spacing',
    'w', 'kern', 'position', 'sz', 'szCs', 'highlight', 'u', 'effect',
    'bdr', 'shd', 'fitText', 'vertAlign', 'rtl', 'cs', 'em', 'lang',
    'eastAsianLayout', 'specVanish', 'oMath'
]

_RPR_ORDER_INDEX = {tag: i for i, tag in enumerate(_RPR_CHILD_ORDER)}


def _reorder_rPr(rPr):
    """Sort all child elements of rPr to strictly match OpenXML ISO/IEC 29500 XSD sequence order.
    Guarantees Microsoft Word correctly respects fonts, colors, bold/italic, size, and RTL without resetting to defaults."""
    if rPr is None or len(rPr) <= 1:
        return

    children = list(rPr)
    for child in children:
        rPr.remove(child)

    def get_order_key(elem):
        tag_name = elem.tag.rsplit('}', 1)[-1] if '}' in elem.tag else elem.tag
        return _RPR_ORDER_INDEX.get(tag_name, 999)

    children.sort(key=get_order_key)
    for child in children:
        rPr.append(child)


def _reorder_pPr(pPr):
    """Sort all child elements of pPr to strictly match OpenXML ISO/IEC 29500 XSD sequence order.
    Guarantees Microsoft Word correctly respects w:bidi, w:spacing, w:ind, w:jc, and w:rPr without discarding alignment."""
    if pPr is None or len(pPr) <= 1:
        return

    children = list(pPr)
    for child in children:
        pPr.remove(child)

    def get_order_key(elem):
        tag_name = elem.tag.rsplit('}', 1)[-1] if '}' in elem.tag else elem.tag
        return _PPR_ORDER_INDEX.get(tag_name, 999)

    children.sort(key=get_order_key)
    for child in children:
        pPr.append(child)


def _apply_paragraph_layout(para, align, effective_rtl):
    """Apply block-level alignment and text direction adhering strictly to OpenXML XSD sequence order.
    Ensures w:bidi and w:jc are placed in correct position so Microsoft Word parses alignment correctly.
    
    IMPORTANT: When w:bidi="1" is present, Word flips w:jc so that "left" aligns visually right,
    and "right" aligns visually left. To counteract this, we map the visual target alignment to the
    appropriate OOXML logical alignment based on the effective text direction.
    """
    pPr = para._p.get_or_add_pPr()

    target_align = None
    if align:
        al = str(align).strip().lower()
        if al in ('right', 'left', 'center', 'justify', 'both', 'start', 'end'):
            target_align = al
        elif al in _OPENXML_JC_MAP:
            target_align = _OPENXML_JC_MAP[al]

    if not target_align:
        target_align = 'right' if effective_rtl else 'left'

    # 1. Set python-docx paragraph.alignment property (using visual target_align)
    align_enum = _PYTHON_DOCX_ALIGN_MAP.get(target_align)
    if align_enum is not None:
        try:
            para.alignment = align_enum
        except Exception:
            pass

    # 2. Determine raw OOXML w:jc value, FLIPPING left/right if effective_rtl is True
    if effective_rtl:
        if target_align == 'left' or target_align == 'end':
            jc_val = 'right'
        elif target_align == 'right' or target_align == 'start':
            jc_val = 'left'
        else:
            jc_val = _OPENXML_JC_MAP.get(target_align, 'left')
    else:
        if target_align == 'start':
            jc_val = 'left'
        elif target_align == 'end':
            jc_val = 'right'
        else:
            jc_val = _OPENXML_JC_MAP.get(target_align, 'left')

    # 3. Set raw OOXML w:jc tag
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), jc_val)

    # 4. Set raw OOXML w:bidi tag (1 for RTL, 0 for LTR)
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '1' if effective_rtl else '0')

    # 5. Sort pPr child elements to strictly comply with OpenXML XSD schema order
    _reorder_pPr(pPr)


def _add_formatted_paragraph(doc, paragraph_text, cat, font_name, font_size,
                              line_spacing, para_indent, space_after, align, effective_rtl, cat_fmt=None, block_font_size=None):
    """Build one docx paragraph from a fragment of HTML/markdown-ish text,
    applying real bold/italic/underline/strike/superscript/subscript/
    color/highlight formatting per run, plus block-level alignment and
    text direction."""
    cat_fmt = cat_fmt or {}
    eff_font_name = cat_fmt.get('fontFamily') or font_name or 'Simplified Arabic'
    
    # Priority hierarchy for font size: block_font_size > cat_fmt['fontSize'] > fallback font_size option > default 14pt
    eff_font_size = None
    if block_font_size:
        m = re.search(r'\d+', str(block_font_size))
        if m: eff_font_size = int(m.group(0))
    if not eff_font_size and cat_fmt.get('fontSize'):
        m = re.search(r'\d+', str(cat_fmt['fontSize']))
        if m: eff_font_size = int(m.group(0))
    if not eff_font_size:
        eff_font_size = font_size or 14

    runs_spec = parse_inline_runs(paragraph_text)
    if not runs_spec:
        return

    # Determine per-paragraph direction (dir="rtl" vs dir="ltr")
    p_dir = _extract_html_dir(paragraph_text)
    for _, style in runs_spec:
        if style.get('dir'):
            p_dir = style['dir']
            break
    if p_dir == 'rtl':
        p_effective_rtl = True
    elif p_dir == 'ltr':
        p_effective_rtl = False
    else:
        p_effective_rtl = _detect_text_direction(paragraph_text, default_rtl=effective_rtl)

    # Determine per-paragraph alignment (style="text-align: ..." or align="...")
    p_align = _extract_html_align(paragraph_text) or align
    for _, style in runs_spec:
        if style.get('align'):
            p_align = style['align']
            break

    if not p_align and cat in ('Title', 'Caption'):
        p_align = 'center'

    p_line_spacing = _extract_html_line_spacing(paragraph_text)
    if p_line_spacing is None and cat_fmt.get('lineSpacing'):
        try: p_line_spacing = float(cat_fmt['lineSpacing'])
        except ValueError: pass
    if p_line_spacing is None:
        p_line_spacing = line_spacing

    p_space_after = _extract_html_margin_bottom(paragraph_text)
    if p_space_after is None and cat_fmt.get('spaceAfter'):
        try: p_space_after = float(cat_fmt['spaceAfter'].replace('pt', '').strip())
        except ValueError: pass
    if p_space_after is None:
        p_space_after = space_after

    p_space_before = _extract_html_margin_top(paragraph_text)
    if p_space_before is None and cat_fmt.get('spaceBefore'):
        try: p_space_before = float(cat_fmt['spaceBefore'].replace('pt', '').strip())
        except ValueError: pass
    if p_space_before is None:
        p_space_before = 0

    para = doc.add_paragraph()

    # Apply paragraph layout (w:bidi & w:jc) BEFORE adding runs
    _apply_paragraph_layout(para, p_align, p_effective_rtl)

    for run_text, style in runs_spec:
        if not run_text:
            continue
        
        # Soft line break (<br>) within the same paragraph
        if run_text == '\n':
            run = para.add_run()
            run.add_break()
            _set_run_font_and_bidi(run, eff_font_name, eff_font_size, is_rtl=p_effective_rtl)
            continue

        run = para.add_run(run_text)
        
        # Run direction: style['dir'] if specified, else paragraph direction
        r_dir = style.get('dir')
        is_r_rtl = (r_dir == 'rtl') if r_dir else (False if r_dir == 'ltr' else p_effective_rtl)

        # Apply comprehensive inline formatting (font, size, CS hint, bold/bCs, italic/iCs, color, highlight)
        _apply_run_formatting(run, style, eff_font_name, eff_font_size, is_r_rtl=is_r_rtl, cat_fmt=cat_fmt)

    if cat in ('Title', 'Section-header'):
        for run in para.runs:
            run.bold = True
            bCs = run._element.get_or_add_rPr().find(qn('w:bCs'))
            if bCs is None:
                bCs = OxmlElement('w:bCs')
                run._element.get_or_add_rPr().append(bCs)
            bCs.set(qn('w:val'), '1')
            _reorder_rPr(run._element.get_or_add_rPr())
    _set_line_spacing(para, p_line_spacing)
    fmt = para.paragraph_format

    # First-line indent should ONLY be applied to 'Text' category when alignment is standard right (or left in LTR).
    std_align = 'right' if p_effective_rtl else 'left'
    normalized_p_align = None
    if p_align:
        al = str(p_align).strip().lower()
        if al == 'start':
            normalized_p_align = 'right' if p_effective_rtl else 'left'
        elif al == 'end':
            normalized_p_align = 'left' if p_effective_rtl else 'right'
        elif al in ('right', 'left', 'center', 'justify', 'both', 'lowkashida', 'mediumkashida', 'highkashida', 'distribute'):
            if al == 'lowkashida': normalized_p_align = 'lowKashida'
            elif al == 'mediumkashida': normalized_p_align = 'mediumKashida'
            elif al == 'highkashida': normalized_p_align = 'highKashida'
            else: normalized_p_align = al

    if cat == 'Text' and (normalized_p_align == std_align or normalized_p_align is None):
        fmt.first_line_indent = Cm(para_indent)
    else:
        fmt.first_line_indent = Cm(0)

    _set_paragraph_spacing(para, p_space_before, p_space_after)

    # Re-apply paragraph layout to guarantee w:bidi and w:jc stay before any paragraph rPr
    _apply_paragraph_layout(para, p_align, p_effective_rtl)


