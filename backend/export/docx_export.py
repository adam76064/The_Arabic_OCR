# docx_export.py - uses shared helpers (including private names)
from .shared import (
    SKIP_CATEGORIES, TEXT_CATEGORIES,
    format_display_text, parse_inline_runs,
    _strip_markdown_and_tags, _split_block_paragraphs,
    _detect_text_direction, _extract_html_dir, _extract_html_align,
    _extract_html_line_spacing, _extract_html_margin_bottom, _extract_html_margin_top,
    _css_color_to_hex, _extract_css_colors, _style_from_tag,
    _set_paragraph_spacing, _set_line_spacing, _add_page_break, _add_page_number_label,
    _set_section_rtl, _set_style_rtl, _set_run_font_and_bidi,
    _set_table_no_borders, _set_row_height, _set_cell_width, _set_cell_valign,
    _apply_poetry_paragraph_layout_v169, _fill_poetry_cell,
    _parse_poetry_lines, _add_poetry_docx,
    _set_run_highlight_hex, _reorder_pPr, _apply_paragraph_layout, _add_formatted_paragraph,
)
# Also bring in constants that import * would normally bring but start without underscore
from .shared import (
    _BLOCK_BOUNDARY_RE, _TAG_RE, _MD_HEADER_RE, _MD_BOLD_RE, _MD_ITALIC_RE,
    _ARABIC_CHAR_RE, _LATIN_CHAR_RE,
    _MD_INLINE_TOKEN_RE, _BOLD_TAGS, _ITALIC_TAGS, _UNDERLINE_TAGS, _SUP_TAGS, _SUB_TAGS,
    _CSS_BOLD_RE, _CSS_ITALIC_RE, _CSS_UNDERLINE_RE, _CSS_STRIKE_RE,
    _CSS_BG_COLOR_RE, _CSS_FG_COLOR_RE, _STRIKE_TAGS, _NAMED_COLORS,
    _PYTHON_DOCX_ALIGN_MAP, _OPENXML_JC_MAP, _PPR_CHILD_ORDER, _PPR_ORDER_INDEX,
    _InlineRunParser,
)
import os, re, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_DIRECTION

import os, re, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_DIRECTION

def export_docx(project, page_indices, output_path, opts=None):
    """
    opts keys (all optional, sensible defaults):
        font_name       str     optional fallback font name
        font_size       int     optional fallback font size
        line_spacing    float   1.0
        para_indent     float   1.0   (cm, first-line indent)
        space_after     int     6     (pt between paragraphs)
        page_numbering  str     'none' | 'pdf' | 'logical'
        page_break      bool    True
        page_size       str     'A4' | 'A5' | 'Letter'
        landscape       bool    False
        rtl             bool    True
        text_mode       str     'formatted' | 'raw'
    """
    opts = opts or {}
    text_mode    = opts.get('text_mode', 'formatted')
    font_name    = opts.get('font_name') or 'Simplified Arabic'
    font_size_val = opts.get('font_size')
    font_size    = int(font_size_val) if font_size_val is not None else 14
    line_spacing = float(opts.get('line_spacing', 1.0))
    para_indent  = float(opts.get('para_indent', 1.0))
    space_after  = int(opts.get('space_after', 6))
    page_numbering = opts.get('page_numbering', 'none')  # none | pdf | logical
    page_break   = bool(opts.get('page_break', True))
    page_size_name = opts.get('page_size', 'A4')
    landscape    = bool(opts.get('landscape', False))
    rtl          = bool(opts.get('rtl', True))
    logical_start = int(project['metadata'].get('logical_start', 1))

    PAGE_SIZES = {
        'A4':     (Cm(21), Cm(29.7)),
        'A5':     (Cm(14.8), Cm(21)),
        'Letter': (Cm(21.59), Cm(27.94)),
    }
    pw, ph = PAGE_SIZES.get(page_size_name, PAGE_SIZES['A4'])
    if landscape:
        pw, ph = ph, pw

    doc = Document()
    section = doc.sections[0]
    section.page_width = pw
    section.page_height = ph
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE

    # 1. Set Section-level BiDi direction
    _set_section_rtl(section, is_rtl=rtl)

    # 2. Set Normal Document Style BiDi & Complex Script font defaults
    style = doc.styles['Normal']
    _set_style_rtl(style, font_name=font_name, font_size_pt=font_size, is_rtl=rtl)

    first_page = True
    for page_num, i in enumerate(page_indices):
        page = project['pages'][i]
        ocr_data = page.get('ocr_data') or []

        if not first_page and page_break:
            _add_page_break(doc)
        first_page = False

        # Page number label at top
        num_font = font_name or 'Simplified Arabic'
        num_size = (font_size - 4) if font_size else 10
        if page_numbering == 'pdf':
            _add_page_number_label(doc, f'[PDF p.{i + 1}]', num_font, num_size, rtl)
        elif page_numbering == 'logical':
            _add_page_number_label(doc, f'[ص {i + logical_start}]', num_font, num_size, rtl)

        cat_fmt_map = (project.get('metadata', {}).get('category_formatting', {}) or {})
        for el in ocr_data:
            cat = el.get('category', 'Text')
            text = el.get('text', '').strip()
            if not text:
                continue
            if cat in SKIP_CATEGORIES:
                continue

            cat_fmt = cat_fmt_map.get(cat, {})
            # NOTE: Do NOT call _extract_html_align(text) here on the full block text.
            # The full text may contain multiple <p> tags with different per-para
            # text-align values — re.search would pick up the FIRST match found anywhere
            # in the string, poisoning the block-level default for all sub-paragraphs.
            # Block-level alignment must come ONLY from the element's stored 'align' field
            # or the category formatting defaults. Per-para HTML alignment is handled
            # inside _add_formatted_paragraph via _extract_html_align(paragraph_text).
            el_align = el.get('align') or cat_fmt.get('align')
            if not el_align and cat in ('Title', 'Caption'):
                el_align = 'center'
            el_dir = el.get('dir') or cat_fmt.get('dir')
            effective_rtl = (el_dir == 'rtl') if el_dir else (False if el_dir == 'ltr' else rtl)
            
            block_font_size = el.get('font_size') or el.get('fontSize') or el.get('size')

            # ── Poetry blocks (Vertical-poetry & Staggered-poetry) ──────────────────────────
            if cat in ('Vertical-poetry', 'Staggered-poetry'):
                _add_poetry_docx(doc, el, cat, font_name, font_size, effective_rtl, cat_fmt=cat_fmt)
                continue

            # ── Generic Table Handling ───────────────────────────────────────────
            if cat == 'Table' and 'table_structure' in el:
                ts = el['table_structure']
                num_rows = ts.get('rows', 1)
                num_cols = ts.get('cols', 1)
                cells = ts.get('cells', [])

                if num_rows > 0 and num_cols > 0:
                    tbl = doc.add_table(rows=num_rows, cols=num_cols)
                    tbl.style = 'Table Grid'
                    if effective_rtl:
                        tbl.table_direction = WD_TABLE_DIRECTION.RTL

                    # 1. Merge cells natively in DOCX based on spans
                    for c_info in cells:
                        r = c_info.get('row', 0)
                        c = c_info.get('col', 0)
                        r_span = c_info.get('row_span', 1)
                        c_span = c_info.get('col_span', 1)

                        if r < num_rows and c < num_cols:
                            top_left_cell = tbl.cell(r, c)
                            if r_span > 1 or c_span > 1:
                                target_r = min(r + r_span - 1, num_rows - 1)
                                target_c = min(c + c_span - 1, num_cols - 1)
                                top_left_cell.merge(tbl.cell(target_r, target_c))

                            # 2. Populate text and preserve formatting!
                            top_left_cell.text = "" # Clear default
                            cell_dir = c_info.get('dir') or el_dir
                            cell_is_rtl = (cell_dir == 'rtl') if cell_dir else (False if cell_dir == 'ltr' else effective_rtl)

                            if text_mode == 'formatted':
                                paragraphs = _split_block_paragraphs(c_info.get('text', '')) or [""]
                                for p_idx, p_text in enumerate(paragraphs):
                                    para = top_left_cell.paragraphs[0] if p_idx == 0 else top_left_cell.add_paragraph()
                                    _set_paragraph_spacing(para, 0, 0)

                                    cell_align = c_info.get('align') or el_align
                                    _apply_paragraph_layout(para, cell_align, cell_is_rtl)

                                    runs_spec = parse_inline_runs(p_text)
                                    for run_text, style in runs_spec:
                                        if not run_text: continue
                                        run = para.add_run(run_text)

                                        cell_r_font = style.get('font_family') or cat_fmt.get('fontFamily') or font_name
                                        cell_cat_sz = None
                                        if cat_fmt.get('fontSize'):
                                            m_sz = re.search(r'\d+', str(cat_fmt['fontSize']))
                                            if m_sz: cell_cat_sz = int(m_sz.group(0))
                                        cell_r_sz = style.get('font_size') or cell_cat_sz or font_size

                                        r_dir = style.get('dir')
                                        is_r_rtl = (r_dir == 'rtl') if r_dir else cell_is_rtl

                                        _set_run_font_and_bidi(run, cell_r_font, cell_r_sz, is_rtl=is_r_rtl)

                                        if style.get('bold'): run.bold = True
                                        if style.get('italic'): run.italic = True
                                        if style.get('underline'): run.underline = True
                                        if style.get('strike'): run.font.strike = True
                                        if style.get('superscript'): run.font.superscript = True
                                        if style.get('subscript'): run.font.subscript = True
                                        if style.get('color'):
                                            try: run.font.color.rgb = RGBColor.from_string(style['color'])
                                            except Exception: pass
                                        if style.get('highlight'):
                                            try: _set_run_highlight_hex(run, style['highlight'])
                                            except Exception: pass
                                    _apply_paragraph_layout(para, cell_align, cell_is_rtl)
                            else:
                                top_left_cell.text = c_info.get('text', '')
                                para = top_left_cell.paragraphs[0]
                                _set_paragraph_spacing(para, 0, 0)
                                _apply_paragraph_layout(para, c_info.get('align') or el_align, cell_is_rtl)
                                for run in para.runs:
                                    _set_run_font_and_bidi(run, font_name, font_size, is_rtl=cell_is_rtl)

                            # 3. Apply background color shading (w:shd)
                            bg_color = c_info.get('bg_color')
                            if bg_color:
                                hex_bg = _css_color_to_hex(bg_color)
                                if hex_bg:
                                    tcPr = top_left_cell._element.get_or_add_tcPr()
                                    shd = OxmlElement('w:shd')
                                    shd.set(qn('w:val'), 'clear')
                                    shd.set(qn('w:color'), 'auto')
                                    shd.set(qn('w:fill'), hex_bg)
                                    tcPr.append(shd)
                continue

            if text_mode == 'formatted':
                paragraphs = _split_block_paragraphs(text) or [text]
                for p_text in paragraphs:
                    _add_formatted_paragraph(doc, p_text, cat, font_name, font_size,
                                              line_spacing, para_indent, space_after,
                                              el_align, effective_rtl, cat_fmt=cat_fmt, block_font_size=block_font_size)
                continue

            # raw mode: keep stored text exactly as-is, one paragraph per element
            para = doc.add_paragraph()
            _apply_paragraph_layout(para, el_align, effective_rtl)
            run = para.add_run(text)
            _set_run_font_and_bidi(run, font_name, block_font_size or font_size, is_rtl=effective_rtl)
            if cat in ('Title', 'Section-header'):
                run.bold = True
            _set_line_spacing(para, line_spacing)
            fmt = para.paragraph_format
            fmt.first_line_indent = Cm(para_indent) if cat == 'Text' else Cm(0)
            _set_paragraph_spacing(para, 0, space_after)
            _apply_paragraph_layout(para, el_align, effective_rtl)

    doc.save(output_path)
    return output_path
