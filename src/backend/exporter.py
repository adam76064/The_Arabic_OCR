"""Export project pages to JSON / TXT / DOCX."""
import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Categories rendered as flowing text (everything else is skipped or handled separately)
TEXT_CATEGORIES = {'Text', 'Section-header', 'Title', 'Caption', 'Footnote', 'List-item'}
SKIP_CATEGORIES = {'Picture', 'Page-header', 'Page-footer', 'Formula'}


def _set_paragraph_spacing(para, space_before=0, space_after=6):
    fmt = para.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)


def _set_line_spacing(para, lines=1.0):
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = lines


def _add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(break_type=1)  # WD_BREAK.PAGE = 1
    _set_paragraph_spacing(para, 0, 0)


def _add_page_number_label(doc, label, font_name, font_size, rtl):
    para = doc.add_paragraph(label)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(para, 0, 0)
    run = para.runs[0]
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.italic = True
    if rtl:
        _set_rtl(para)


def _set_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def export_json(project, page_indices, output_path):
    pages = [project['pages'][i] for i in page_indices]
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump({'project_id': project['id'],
                   'metadata': project['metadata'],
                   'pages': pages}, f, ensure_ascii=False, indent=2)
    return output_path


def export_txt(project, page_indices, output_path, logical_start=1):
    lines = []
    for i in page_indices:
        page = project['pages'][i]
        for el in (page.get('ocr_data') or []):
            if el.get('category') in SKIP_CATEGORIES:
                continue
            lines.append(el.get('text', '').strip())
        lines.append(f"\n— صفحة {i + logical_start} —\n")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    return output_path


def export_docx(project, page_indices, output_path, opts=None):
    """
    opts keys (all optional, sensible defaults):
        font_name       str     'Simplified Arabic'
        font_size       int     16
        line_spacing    float   1.0
        para_indent     float   1.0   (cm, first-line indent)
        space_after     int     6     (pt between paragraphs)
        page_numbering  str     'none' | 'pdf' | 'logical'
        page_break      bool    True
        page_size       str     'A4' | 'A5' | 'Letter'
        landscape       bool    False
        rtl             bool    True
    """
    opts = opts or {}
    font_name    = opts.get('font_name', 'Simplified Arabic')
    font_size    = int(opts.get('font_size', 16))
    line_spacing = float(opts.get('line_spacing', 1.0))
    para_indent  = float(opts.get('para_indent', 1.0))
    space_after  = int(opts.get('space_after', 6))
    page_numbering = opts.get('page_numbering', 'none')  # none | pdf | logical
    page_break   = bool(opts.get('page_break', True))
    page_size_name = opts.get('page_size', 'A4')
    landscape    = bool(opts.get('landscape', False))
    rtl          = bool(opts.get('rtl', True))
    logical_start = int(project['metadata'].get('logical_start', 1))

    PAGE_SIZES = {
        'A4':     (Cm(21), Cm(29.7)),
        'A5':     (Cm(14.8), Cm(21)),
        'Letter': (Cm(21.59), Cm(27.94)),
    }
    pw, ph = PAGE_SIZES.get(page_size_name, PAGE_SIZES['A4'])
    if landscape:
        pw, ph = ph, pw

    doc = Document()
    section = doc.sections[0]
    section.page_width = pw
    section.page_height = ph
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE

    # Default style
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)

    first_page = True
    for page_num, i in enumerate(page_indices):
        page = project['pages'][i]
        ocr_data = page.get('ocr_data') or []

        if not first_page and page_break:
            _add_page_break(doc)
        first_page = False

        # Page number label at top
        if page_numbering == 'pdf':
            _add_page_number_label(doc, f'[PDF p.{i + 1}]', font_name, font_size - 4, rtl)
        elif page_numbering == 'logical':
            _add_page_number_label(doc, f'[ص {i + logical_start}]', font_name, font_size - 4, rtl)

        for el in ocr_data:
            cat = el.get('category', 'Text')
            text = el.get('text', '').strip()
            if not text:
                continue
            if cat in SKIP_CATEGORIES:
                continue

            if cat == 'Table':
                # Simple table: split lines into rows, first line = header
                rows = [r.split('|') for r in text.splitlines() if r.strip()]
                rows = [r for r in rows if any(c.strip() for c in r)]
                if rows:
                    cols = max(len(r) for r in rows)
                    tbl = doc.add_table(rows=len(rows), cols=cols)
                    tbl.style = 'Table Grid'
                    for ri, row in enumerate(rows):
                        for ci, cell_text in enumerate(row[:cols]):
                            cell = tbl.rows[ri].cells[ci]
                            cell.text = cell_text.strip()
                            for para in cell.paragraphs:
                                _set_paragraph_spacing(para, 0, 0)
                                run = para.runs[0] if para.runs else para.add_run(cell.text)
                                run.font.name = font_name
                                run.font.size = Pt(font_size)
                continue

            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            if cat in ('Title', 'Section-header'):
                run.bold = True
            _set_line_spacing(para, line_spacing)
            fmt = para.paragraph_format
            fmt.first_line_indent = Cm(para_indent) if cat == 'Text' else Cm(0)
            _set_paragraph_spacing(para, 0, space_after)
            if rtl:
                _set_rtl(para)

    doc.save(output_path)
    return output_path
